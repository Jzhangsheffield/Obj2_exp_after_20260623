#!/usr/bin/env python
"""Create the Chinese Markdown and DOCX experiment report from analysis tables."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "analysis" / "obj2_experiment_report_20260719"
TABLES = OUT / "tables"
FIGURES = OUT / "figures"
REPORT_MD = OUT / "Obj2_完整实验分析报告_扩充版_20260720.md"
REPORT_DOCX = OUT / "Obj2_完整实验分析报告_扩充版_20260720.docx"

NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREY = "F2F4F7"
MID_GREY = "666666"
GREEN = "2E7D54"
RED = "A23B3B"
GOLD = "9A6A18"


def load(name: str) -> pd.DataFrame:
    return pd.read_csv(TABLES / name)


selected = load("selected_best_available.csv")
strict_last = load("strict_last_checkpoint.csv")
family = load("family_summary.csv")
ft = load("ft_training_runs.csv")
cl = load("cl_configs_deduplicated.csv")
issues = load("quality_issues.csv")
diagnostics = load("model_diagnostics.csv")
confusions = load("confusion_pairs.csv")
feature = load("feature_separation_pilot.csv")
module_summary = load("module_method_summary.csv")
selected_module_models = load("selected_module_models.csv")
selected_module_per_class = load("selected_module_per_class.csv")
module_bootstrap = load("paired_bootstrap_summary.csv")
class_effect_summary = load("class_effect_summary.csv")
feature_diag = load("targeted_feature_diagnostics.csv")
harmed_selection = load("harmed_class_selection.csv")
metric_defs = load("metric_definitions.csv")
figure_guide = load("figure_reading_guide.csv")
figure_guide.loc[figure_guide.figure_type == "Shared UMAP feature plots", "figure_type"] = "Shared aligned PCA feature plots"
figure_guide.loc[figure_guide.figure_type == "Shared aligned PCA feature plots", "encoding"] = "The same samples are projected to a common 30-D basis, orthogonally aligned to SupCon, and jointly reduced to 2-D PCA. Circles are harmed classes; x marks are common confusion targets."
figure_guide.loc[figure_guide.figure_type == "Shared aligned PCA feature plots", "how_to_read"] = "Compare mixing and compactness, not absolute coordinate meaning; confirm with high-dimensional silhouette, purity, and distance metrics."
summary = json.loads((OUT / "analysis_summary.json").read_text(encoding="utf-8"))


def f3(value) -> str:
    return "-" if pd.isna(value) else f"{float(value):.3f}"


def pp(value) -> str:
    return "-" if pd.isna(value) else f"{float(value)*100:+.1f} pp"


def short(text: str, n: int = 52) -> str:
    text = str(text)
    return text if len(text) <= n else text[:n-1] + "…"


def best_pretrained(scope: str, modality: str, mode: str, frame=selected):
    d = frame[(frame.dataset_scope == scope) & (frame.modality == modality) &
              (frame.finetune_mode == mode) & (frame.finetune_mode != "scratch")]
    return None if d.empty else d.sort_values("test_balanced_acc", ascending=False).iloc[0]


def best_scratch(scope: str, modality: str, frame=selected):
    d = frame[(frame.dataset_scope == scope) & (frame.modality == modality) & (frame.finetune_mode == "scratch")]
    return None if d.empty else d.sort_values("test_balanced_acc", ascending=False).iloc[0]


def checkpoint_tag(row) -> str:
    return "best_val" if row is not None and row.checkpoint_selection == "best_val" else "last"


def mean_by_run(family_name: str) -> pd.DataFrame:
    d = selected[selected.family == family_name].copy()
    return (d.groupby(["run_dir", "source_key", "finetune_mode"], as_index=False)
            .agg(mean_balanced=("test_balanced_acc", "mean"), min_balanced=("test_balanced_acc", "min"),
                 max_balanced=("test_balanced_acc", "max"), n_views=("test_balanced_acc", "count"),
                 backbone_lr=("cfg_backbone_learning_rate", "first"))
            .sort_values("mean_balanced", ascending=False))


except_rows = []
for modality in ("rgb", "emg", "imu"):
    scratch = best_scratch("except_take_put", modality)
    full = best_pretrained("except_take_put", modality, "full")
    head = best_pretrained("except_take_put", modality, "head_only")
    except_rows.append((modality.upper(), scratch, full, head))

take_rows = []
for modality in ("rgb", "emg", "imu"):
    scratch = best_scratch("take_put", modality)
    full = best_pretrained("take_put", modality, "full")
    head = best_pretrained("take_put", modality, "head_only")
    take_rows.append((modality.upper(), scratch, full, head))

round2 = mean_by_run("round2")
dualcam = mean_by_run("dualcam")
round2_best = round2[round2.finetune_mode != "scratch"].iloc[0]
round2_scratch = round2[round2.finetune_mode == "scratch"].iloc[0]
dualcam_best = dualcam[dualcam.finetune_mode != "scratch"].iloc[0]
dualcam_scratch = dualcam[dualcam.finetune_mode == "scratch"].iloc[0]

overfit = (ft.groupby(["dataset_scope", "modality", "finetune_mode"], as_index=False)
           .agg(n=("model_id", "count"), best_epoch_median=("best_val_balanced_epoch", "median"),
                train_val_gap_median=("train_val_acc_gap", "median"),
                best_final_drop_median=("best_to_final_balanced_drop", "median")))
cal = (diagnostics.groupby("modality", as_index=False)
       .agg(n_models=("model_id", "count"), mean_ece=("ece_10bin", "mean"),
            mean_confidence=("mean_confidence", "mean"), error_rate=("error_rate", "mean"),
            high_conf_errors=("high_confidence_errors", "sum")))

MODULE_NAME = {
    "prototype": "SupCon+Proto",
    "relation": "SupCon+Rel",
    "prototype+relation": "SupCon+Proto+Rel",
}


def selected_module_rows():
    rows = []
    for modality in ("rgb", "emg", "imu"):
        m = selected_module_models[(selected_module_models.modality == modality) &
                                   (selected_module_models.model_role == "module")].iloc[0]
        bs = module_bootstrap[(module_bootstrap.modality == modality) &
                              (module_bootstrap.comparison == "module_vs_scratch")].iloc[0]
        bu = module_bootstrap[(module_bootstrap.modality == modality) &
                              (module_bootstrap.comparison == "module_vs_supcon")].iloc[0]
        rows.append((modality.upper(), m, bs, bu))
    return rows


FAMILY_DESCRIPTIONS = [
    ("基础损失组合", "验证 SupCon、prototype、relation 及其组合", "原型数 1/2/3；EMA 0.3/0.5/0.8；22 个损失配置/模态", "except_take_put + take_put"),
    ("队列容量 K 扫描", "检验负样本队列容量", "K=512/1024/2048（MindRove 对应 1024/2048 等）；SupCon 与 proto+relation", "except_take_put"),
    ("关系损失 Top-k", "限制不同类关系损失参与范围", "Top-k=5/10/all；固定代表性 prototype/relation 配置", "except_take_put"),
    ("平衡批采样", "改善类别不平衡", "balanced_batch；每类样本数/每批类别数；与 random sampler 对照", "except_take_put"),
    ("阶段延迟 + Top-k", "提前启用 prototype/relation 并扫描 Top-k", "loss start epoch: 50→5；Top-k=3/5/10/all", "except_take_put"),
    ("RGB ResNet-10", "检验更浅 RGB backbone", "model_depth: 18→10；4 个代表性损失配置", "except_take_put"),
    ("RGB 动作语义保持增强", "减少破坏动作方向/空间语义的增强", "弱增强、shared/overlap75、relation/CE、backbone LR 1e-4/3e-4", "Round-2"),
    ("RGB 双相机预训练", "检验同相机、交叉相机、混合相机预训练", "same143/same152/cross/hybrid；Cam 143 与 152 分别测试", "Dual-camera"),
]


def md_table(headers: list[str], rows: Iterable[Iterable[str]]) -> str:
    lines = ["| " + " | ".join(headers) + " |", "|" + "|".join(["---"]*len(headers)) + "|"]
    lines.extend("| " + " | ".join(str(x).replace("|", "\\|") for x in row) + " |" for row in rows)
    return "\n".join(lines)


def build_markdown() -> str:
    inv = summary["inventory"]
    lines = [
        "# Obj2 多模态对比学习与微调实验：完整分析报告",
        "",
        "生成日期：2026-07-20  ",
        f"数据快照：{summary['generated_at']}  ",
        "主要指标：测试集 balanced accuracy；同时报告 accuracy、macro-F1、类别 recall、校准与混淆。",
        "",
        "## 执行摘要",
        "",
        f"本报告汇总 {inv['ft_training_runs']} 个微调运行、{inv['test_metric_rows']} 条测试汇总记录，以及去重后的 {inv['cl_configs_deduplicated']} 个对比学习配置。",
        "",
        "- **IMU 是当前最稳定受益于预训练的模态。** except_take_put 的最佳 head-only balanced accuracy 为 0.604，较对应 scratch 高约 7.9 pp；take_put 的最佳值为 0.666，较 scratch 高 4.5 pp。",
        "- **EMG 有中等收益，但绝对性能和校准仍弱。** except_take_put 最佳 head-only 为 0.369，较 scratch 高 6.4 pp；平均 ECE 约 0.388。",
        "- **传统 RGB 预训练尚未稳定超过 scratch。** except_take_put 最佳 full 预训练为 0.607，而跨实验族最佳 scratch 达 0.638；head-only 更明显落后，说明 RGB 需要更新 backbone。",
        f"- **Round-2 的动作语义保持增强是最有价值的 RGB 方向。** `{round2_best.source_key}`、backbone LR={round2_best.backbone_lr:g} 的 balanced accuracy 达 {round2_best.max_balanced:.3f}，相对该轮 scratch 提升 {(round2_best.max_balanced-round2_scratch.mean_balanced)*100:.1f} pp。",
        f"- **双相机预训练当前失败。** 最佳预训练配置跨相机平均 {dualcam_best.mean_balanced:.3f}，低于 scratch 的 {dualcam_scratch.mean_balanced:.3f}。",
        "- **结果可信度的首要限制是仅有 seed=1，以及 RGB scratch 跨实验族波动较大。** 在补齐多 seed 和统一 checkpoint 选择前，不应把 1-2 pp 的差异视为确定收益。",
        "",
        "## 1. 数据覆盖与分析规则",
        "",
        f"- 原始 args.json：{inv['cl_args_files_raw']} 个；去除 RGB `cl_`/非 `cl_` 重复目录后为 {inv['cl_configs_deduplicated']} 个配置。",
        f"- 微调 summary.json：{inv['ft_training_runs']} 个。测试记录中 best_val={inv['best_val_test_rows']}，last={inv['last_test_rows']}。",
        f"- 最佳可用视图包含 {inv['selected_models_or_camera_views']} 个“运行/相机视图”，全部可回溯到逐样本 CSV。",
        f"- 分析统计与绘图已在 `{summary.get('analysis_runtime', {}).get('python_executable', 'unknown')}` 环境中重新运行。",
        "- 主视图优先使用 best_val；如果该实验没有 best_val 测试，则使用 last。报告同时保留严格 last 表，用于同 checkpoint 比较。",
        "- 同一实验族内部使用其 scratch 作为主要参考；跨实验族比较时单独标注 checkpoint 与 scratch 波动。",
        "",
        "## 2. 实验设计总览",
        "",
        md_table(["实验族", "研究问题", "主要变量", "范围"], FAMILY_DESCRIPTIONS),
        "",
        "## 3. except_take_put 主任务结果",
        "",
        md_table(["模态", "最佳 scratch", "最佳 full", "最佳 head-only", "判断"], [
            [m, f"{f3(s.test_balanced_acc)} ({checkpoint_tag(s)})",
             f"{f3(f.test_balanced_acc)} / {short(f.source_key,35)} ({checkpoint_tag(f)})",
             f"{f3(h.test_balanced_acc)} / {short(h.source_key,35)} ({checkpoint_tag(h)})",
             ("预训练未稳定超过强 scratch；必须 full" if m=="RGB" else "预训练有收益；head-only 略优" if m=="EMG" else "预训练收益最稳定；full/head 接近")]
            for m,s,f,h in except_rows]),
        "",
        "![实验族性能对比](figures/family_best_balanced_accuracy.png)",
        "",
        "### RGB",
        "",
        "基础损失组合中，`suploss_proto_p1` 的 full 微调 balanced accuracy 为 0.607，比该实验族的 scratch 高 2.3 pp；但其他新实验族产生的 scratch 可达 0.638。因此当前证据说明 prototype 预训练可能有局部收益，但尚未形成可复现的跨运行优势。head-only 的最佳值只有 0.475，说明冻结 RGB backbone 会保留不适合测试域的表示。",
        "",
        "### EMG 与 IMU",
        "",
        "EMG 的最佳配置来自 relation Top-k 的 head-only（0.369）；IMU 的最佳 head-only 来自 K_queue 扫描（0.604）。IMU 对 prototype+relation 的收益比 EMG 更一致；EMG 的不同配置差距较小，且高置信错误明显更多。",
        "",
        "## 4. take_put 二分类结果",
        "",
        md_table(["模态", "scratch", "最佳 full", "相对 scratch", "最佳 head-only", "相对 scratch"], [
            [m, f3(s.test_balanced_acc), f3(f.test_balanced_acc), pp(f.test_balanced_acc-s.test_balanced_acc),
             f3(h.test_balanced_acc), pp(h.test_balanced_acc-s.test_balanced_acc)] for m,s,f,h in take_rows]),
        "",
        "RGB 仅 full 微调略高于 scratch，head-only 明显退化；IMU/EMG 的 full 与 head-only 均优于 scratch。这与 except_take_put 的模态规律一致：RGB 需要更新 backbone，而一维信号预训练更适合冻结或小幅更新。",
        "",
        "## 5. RGB 后续实验：Round-2 与双相机",
        "",
        f"Round-2 最佳单次结果为 `{round2_best.source_key}`，backbone LR={round2_best.backbone_lr:g}，balanced accuracy={round2_best.max_balanced:.3f}。同一配置在较低 backbone LR 下明显较弱，说明学习率和增强策略存在交互。",
        f"双相机实验中，最佳预训练配置 `{dualcam_best.source_key}` 跨 Cam 143/152 平均仅 {dualcam_best.mean_balanced:.3f}，而 scratch 为 {dualcam_scratch.mean_balanced:.3f}。当前 cross-camera 预训练没有带来域泛化，部分配置在两个相机上都退化。",
        "",
        "![RGB 后续实验](figures/round2_and_dualcam.png)",
        "",
        "## 6. 参数敏感性与 checkpoint",
        "",
        "K_queue 对三种模态的平均影响都小于模态本身差异：RGB 在较大 K 下略升，IMU 在中间值较好，EMG 效果较弱。Top-k 的最优值也具有模态依赖，未出现跨模态统一最优点。",
        "",
        "![参数敏感性](figures/hyperparameter_sensitivity.png)",
        "",
        "best_val 相对 last 的平均变化多数在 ±1 pp 内，但单个运行可超过 4 pp。因此 checkpoint 不一致不是全部性能差异的来源，却足以改变接近配置的排序。",
        "",
        "![checkpoint 影响](figures/checkpoint_selection_effect.png)",
        "",
        "## 7. 训练稳定性、校准与类别问题",
        "",
        "except_take_put 的训练-验证 accuracy gap 中位数：EMG full 约 0.238、IMU full 约 0.144、RGB full 约 0.112。IMU 与 RGB 的 head-only 通常更早达到最佳验证 epoch；EMG full/head 均在约 45 epoch 达峰。",
        "",
        md_table(["模态", "模型/视图数", "平均 ECE", "平均置信度", "平均错误率", "高置信错误*"], [
            [r.modality.upper(), str(int(r.n_models)), f3(r.mean_ece), f3(r.mean_confidence), f3(r.error_rate), str(int(r.high_conf_errors))]
            for _,r in cal.iterrows()]),
        "",
        "\\* 高置信错误为模型×样本累计次数，不是唯一原始样本数。EMG 最严重，适合优先尝试 temperature scaling，并检查信号归一化和类别边界。",
        "",
        "![类别 recall 变化](figures/class_recall_delta_vs_scratch.png)",
        "",
        "常见错误包括 EMG 的 wrap→pull_out、insert→pull_out，RGB 的 pull_out→insert、insert→label/wrap，以及 take_put 中 put→take。类别级热力图还显示 RGB 预训练对 press 有明显帮助，但对 measure、label、insert 等类别可能退化。",
        "",
        "## 8. 特征分布诊断（复用已有关键特征）",
        "",
        "本轮没有重新遍历全部权重。复用了旧分析已经提取的代表性 512D 特征：RGB head-only 测试 silhouette=-0.229，而 full=0.085；EMG 各测试表示仍接近或低于 0；IMU full 的测试 silhouette 最高约 0.089。训练集与测试集分离度差距很大，说明域偏移仍是核心问题。",
        "",
        "![特征分离度](figures/feature_separation_train_test.png)",
        "",
        "RGB head-only UMAP 呈现由采集条件主导的轨迹结构、类别颜色严重混合；full 微调后局部类别簇更清晰，但仍存在大量重叠。这与性能结果一致：RGB 不能只训练分类头。",
        "",
        "![RGB head-only UMAP](figures/rgb_head_test_umap.png)",
        "",
        "![RGB full UMAP](figures/rgb_full_test_umap.png)",
        "",
        "## 9. 主要问题与下一步",
        "",
    ]
    for _,r in issues.iterrows():
        lines += [f"### [{r.severity}] {r.issue}", "", str(r.evidence), "", f"影响：{r.impact}", "", f"建议：{r.recommended_action}", ""]
    lines += [
        "## 10. 推荐决策",
        "",
        "1. 将 IMU 的 K_queue/prototype+relation head-only 与基础 `suploss_proto_p2` full 作为首批多 seed 候选。",
        "2. 将 EMG 的 relation Top-k all head-only 作为候选，同时加入归一化与校准消融。",
        "3. RGB 以 Round-2 `a2_weak_shared_supcon`、backbone LR=3e-4 为主线，和统一 scratch、基础 `proto_p1` full 做 5-seed 对照。",
        "4. 暂停扩展双相机预训练网格；先定位 cross-camera 表示退化原因，检查相机归一化、采样对齐和预训练/微调域的一致性。",
        "5. 统一保存 best_val、last、代码 commit、数据 manifest 哈希和随机性设置，再继续细粒度超参数搜索。",
        "",
        "## 附录：指标与可比性",
        "",
        "- balanced accuracy：各类别 recall 的宏平均，是本报告主排序指标。",
        "- macro-F1：逐类别 F1 的宏平均，用于补充不平衡类别评估。",
        "- ECE：10-bin expected calibration error，越低越好。",
        "- best-available：优先 best_val，无对应测试时退回 last；严格 last 明细保存在工作簿和 CSV。",
        "- 本报告中的显著提升均是描述性差值，不是统计显著性结论；当前仅 seed=1。",
    ]
    lines += [
        "",
        "## 专项扩充：Prototype loss 与 Relation loss",
        "",
        "这里把 scratch（从头训练）和 SupCon-only 都视为基线，并在同一实验族、同一模态和同一微调模式内比较 SupCon+Proto、SupCon+Rel、SupCon+Proto+Rel。热力图单元格是该组合下观测到的最佳 balanced accuracy；空白表示没有对应实验，不能按 0 解释。",
        "",
        md_table(["模态", "最佳模块", "实验簇/模式", "Bal. Acc", "Δ Scratch（95% CI）", "Δ SupCon（95% CI）"], [
            [mod, MODULE_NAME.get(m.loss_group, m.loss_group), f"{m.family}/{m.finetune_mode}", f3(m.test_balanced_acc),
             f"{pp(bs.delta_balanced_acc)} [{pp(bs.ci_low)}, {pp(bs.ci_high)}]",
             f"{pp(bu.delta_balanced_acc)} [{pp(bu.ci_low)}, {pp(bu.ci_high)}]"]
            for mod,m,bs,bu in selected_module_rows()
        ]),
        "",
        "RGB 的 SupCon+Proto 相对同簇 SupCon 提升约 9.1 pp，但相对 scratch 仅约 2.3 pp，且样本 bootstrap 区间跨 0；因此它说明 Proto 修复了较弱 SupCon 表示的一部分问题，却还不能证明优于强 scratch。EMG 的最佳模块是 SupCon+Rel，IMU 的最佳模块是 K=2048 下的 SupCon+Proto+Rel；二者相对两种基线的区间均在 0 以上，但区间只反映固定权重上的测试样本不确定性，不包括训练 seed 波动。",
        "",
    ]
    for modality in ("rgb", "emg", "imu"):
        lines += [f"![{modality.upper()} 模块比较](figures/module_method_performance_{modality}.png)", "",
                  "读图：逐行比较同一实验簇；列对应 Scratch、SupCon、SupCon+Proto、SupCon+Rel、SupCon+Proto+Rel，颜色越深表示 balanced accuracy 越高。不同实验簇的 checkpoint 和参数空间可能不同，因此优先做行内比较。", ""]
    lines += [
        "### 最佳模块的逐类别 recall 变化",
        "",
        "每个模态选择一个最佳自研模块，并与匹配的 Scratch 和 SupCon 比较。上半图给绝对 recall，下半图给模块减基线的变化；误差线是配对样本 bootstrap 95% 区间，虚线 ±5 pp 是描述性筛选阈值，不是显著性阈值。",
        "",
    ]
    for modality in ("rgb", "emg", "imu"):
        lines += [f"![{modality.upper()} 逐类别 recall](figures/selected_module_per_class_{modality}.png)", ""]
    lines += [
        "### 每种方法对哪些类别起作用或产生副作用",
        "",
        "以下热力图在基础损失实验簇内，为每种方法选取 balanced accuracy 最佳的配置；绿色为 recall 增加、红色为下降。按 ±5 pp 统计时，RGB full 的 Proto 相对 SupCon 对 9 类有益、3 类有害；EMG head-only 的 Rel 对 7 类有益、1 类有害；IMU head-only 的 Proto+Rel 对 5 类有益、4 类有害。完整类别名称、幅度与所选 source_key 在工作簿 Class Effects / Effect Summary 中。",
        "",
    ]
    for modality in ("rgb", "emg", "imu"):
        lines += [f"![{modality.upper()} 类别作用矩阵](figures/class_effects_{modality}.png)", "",
                  "读图：每一行是一种损失方法，每一列是类别；数值为 recall 相对面板标题所示基线的变化。接近白色表示变化较小，不能把颜色强弱直接解释为跨类别统计显著性。", ""]
    lines += [
        "### 退化类别的定向特征空间对照",
        "",
        "仅对最佳模块相对 SupCon 确实下降的类别及其最常见混淆目标提取特征；三种模型使用同一批测试样本，先投影到共同 30D 子空间并正交对齐到 SupCon，再联合做二维 PCA。圆点是退化类别，叉号是混淆目标。坐标轴没有物理单位，应比较簇的混合、紧凑度和相对邻近关系，并结合下表的原始高维指标判断。",
        "",
        md_table(["模态", "模型", "样本", "cosine silhouette", "5-NN purity", "类内 cosine 距离", "最近类中心距离", "中心 margin"], [
            [r.modality.upper(), r.role, str(int(r.n_samples)), f3(r.silhouette_cosine), f3(r.knn5_purity),
             f3(r.within_class_cosine_distance), f3(r.nearest_other_centroid_distance), f3(r.centroid_margin)]
            for _,r in feature_diag[feature_diag.scope == "selected_harmed_and_confusion_classes"].iterrows()
        ]),
        "",
        "高维诊断显示：RGB 模块相对 SupCon 的 silhouette 略升（约 0.136→0.158）、类内距离下降，但 5-NN purity 基本不变，说明整体几何略改善却仍可在 tear、insert、label 等局部决策边界上退化。EMG 的 Rel 模块 silhouette 略升，但类内距离也增大，pull_out 的副作用更像局部类边界交换。IMU 模块相对 SupCon 的 silhouette、purity 和最近类中心距离均下降，与 open recall 下降一致；这提示 Proto+Rel 的总体收益并非所有类别共享。",
        "",
    ]
    for modality in ("rgb", "emg", "imu"):
        lines += [f"![{modality.upper()} 退化类别共享 PCA](figures/harmed_class_feature_space_{modality}.png)", ""]
    lines += [
        "### 校准图与指标计算说明",
        "",
        "可靠性图把预测置信度划分为 10 个区间。横轴是区间平均置信度，纵轴是实际正确率；曲线落在对角线下方表示过度自信。ECE = Σ_b (n_b/N)·|acc(b)−conf(b)|，因此同时受偏差大小和该区间样本占比影响。",
        "",
        "![所选模型可靠性图](figures/selected_model_reliability.png)",
        "",
        md_table(["指标", "计算", "作用", "解释限制"], metric_defs[["metric", "calculation", "purpose", "interpretation_limit"]].values.tolist()),
        "",
        "### 图表阅读指南",
        "",
        md_table(["图表", "编码含义", "阅读方式"], figure_guide[["figure_type", "encoding", "how_to_read"]].values.tolist()),
    ]
    return "\n".join(lines) + "\n"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar"); tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None: node = OxmlElement(f"w:{m}"); tc_mar.append(node)
        node.set(qn("w:w"), str(v)); node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]):
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None: tbl_w = OxmlElement("w:tblW"); tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa))); tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None: tbl_ind = OxmlElement("w:tblInd"); tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120"); tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr(); tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None: tc_w = OxmlElement("w:tcW"); tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx])); tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr(); node = OxmlElement("w:tblHeader"); node.set(qn("w:val"), "true"); tr_pr.append(node)


def set_cant_split(row):
    """Keep a table row on one page so a checkpoint label cannot spill alone."""
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_run_font(run, name="Calibri", size=11, bold=None, color=None, italic=None):
    run.font.name = name; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name); run._element.rPr.rFonts.set(qn("w:hAnsi"), name); run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor.from_string(color)


def add_numbering(doc: Document, kind: str) -> int:
    root = doc.part.numbering_part.element
    abs_ids = [int(x.get(qn("w:abstractNumId"))) for x in root.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in root.findall(qn("w:num"))]
    aid, nid = (max(abs_ids)+1 if abs_ids else 1), (max(num_ids)+1 if num_ids else 1)
    abstract = OxmlElement("w:abstractNum"); abstract.set(qn("w:abstractNumId"), str(aid))
    mlt = OxmlElement("w:multiLevelType"); mlt.set(qn("w:val"), "singleLevel"); abstract.append(mlt)
    lvl = OxmlElement("w:lvl"); lvl.set(qn("w:ilvl"), "0"); abstract.append(lvl)
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
    fmt = OxmlElement("w:numFmt"); fmt.set(qn("w:val"), "bullet" if kind=="bullet" else "decimal"); lvl.append(fmt)
    txt = OxmlElement("w:lvlText"); txt.set(qn("w:val"), "•" if kind=="bullet" else "%1."); lvl.append(txt)
    jc = OxmlElement("w:lvlJc"); jc.set(qn("w:val"), "left"); lvl.append(jc)
    ppr = OxmlElement("w:pPr"); tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "720"); tabs.append(tab); ppr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "720"); ind.set(qn("w:hanging"), "360"); ppr.append(ind); lvl.append(ppr)
    root.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(nid)); ref = OxmlElement("w:abstractNumId"); ref.set(qn("w:val"), str(aid)); num.append(ref); root.append(num)
    return nid


def set_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr(); numpr = OxmlElement("w:numPr"); ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0"); nid = OxmlElement("w:numId"); nid.set(qn("w:val"), str(num_id)); numpr.extend([ilvl,nid]); ppr.append(numpr)


def add_bullet(doc, text, num_id):
    p = doc.add_paragraph(); set_num(p,num_id); p.paragraph_format.space_after=Pt(8); p.paragraph_format.line_spacing=1.167; set_run_font(p.add_run(text)); return p


def add_number(doc, text, num_id):
    p = doc.add_paragraph(); set_num(p,num_id); p.paragraph_format.space_after=Pt(8); p.paragraph_format.line_spacing=1.167; set_run_font(p.add_run(text)); return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers)); table.style = "Table Grid"
    for i,h in enumerate(headers):
        cell=table.rows[0].cells[i]; cell.text=""; set_cell_shading(cell,LIGHT_GREY); p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_run_font(p.add_run(str(h)),size=9,bold=True,color=NAVY)
    set_repeat_header(table.rows[0]); set_cant_split(table.rows[0])
    for row in rows:
        table_row=table.add_row(); set_cant_split(table_row); cells=table_row.cells
        for i,value in enumerate(row):
            cells[i].text=""; p=cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.LEFT if i==len(row)-1 else WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); set_run_font(p.add_run(str(value)),size=8.5)
    set_table_geometry(table,widths); doc.add_paragraph().paragraph_format.space_after=Pt(2); return table


def add_caption(doc, text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(8); set_run_font(p.add_run(text),size=9,italic=True,color=MID_GREY)


def add_picture(doc, name, caption, width=6.25):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True
    shape=p.add_run().add_picture(str(FIGURES/name),width=Inches(width))
    shape._inline.docPr.set("descr", str(caption))
    shape._inline.docPr.set("title", str(name))
    add_caption(doc,caption)


def add_page_field(paragraph):
    paragraph.alignment=WD_ALIGN_PARAGRAPH.RIGHT; run=paragraph.add_run(); begin=OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"),"begin"); instr=OxmlElement("w:instrText"); instr.set(qn("xml:space"),"preserve"); instr.text=" PAGE "; end=OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"),"end"); run._r.extend([begin,instr,end])


def build_docx():
    doc=Document(); sec=doc.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11); sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1); sec.header_distance=sec.footer_distance=Inches(.492)
    styles=doc.styles
    normal=styles["Normal"]; normal.font.name="Calibri"; normal.font.size=Pt(11); normal._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); normal.paragraph_format.space_before=Pt(0); normal.paragraph_format.space_after=Pt(6); normal.paragraph_format.line_spacing=1.10
    for name,size,color,before,after in [("Title",28,NAVY,0,8),("Subtitle",14,MID_GREY,0,8),("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK_BLUE,8,4)]:
        st=styles[name]; st.font.name="Calibri"; st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=name!="Subtitle"; st._element.rPr.rFonts.set(qn("w:eastAsia"),"Microsoft YaHei"); st.paragraph_format.space_before=Pt(before); st.paragraph_format.space_after=Pt(after); st.paragraph_format.keep_with_next=True
    # A clean, empty header avoids alternating-page artifacts in Word/PDF export.
    sec.header.paragraphs[0].text = ""
    add_page_field(sec.footer.paragraphs[0])
    bullet_id=add_numbering(doc,"bullet"); number_id=add_numbering(doc,"decimal")

    doc.add_paragraph().paragraph_format.space_after=Pt(92)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_run_font(p.add_run("EXPERIMENT REPORT"),size=10,bold=True,color=GOLD)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(8); set_run_font(p.add_run("Obj2 多模态对比学习与微调实验"),size=28,bold=True,color=NAVY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_run_font(p.add_run("完整分析报告"),size=20,bold=True,color=BLUE)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(28); set_run_font(p.add_run("RGB · EMG · IMU | 对比学习 · 微调 · 类别诊断"),size=11,color=MID_GREY)
    doc.add_paragraph().paragraph_format.space_after=Pt(120)
    for text in ["生成日期：2026-07-20",f"数据快照：{summary['generated_at']}","主指标：测试集 balanced accuracy"]:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_run_font(p.add_run(text),size=10,color=MID_GREY)
    doc.add_page_break()

    doc.add_heading("执行摘要",level=1)
    inv=summary["inventory"]
    p=doc.add_paragraph(); set_run_font(p.add_run("分析范围："),bold=True,color=NAVY); set_run_font(p.add_run(f"{inv['ft_training_runs']} 个微调运行、{inv['test_metric_rows']} 条测试记录、去重后的 {inv['cl_configs_deduplicated']} 个对比学习配置。"))
    bullets=[
        "IMU 是最稳定受益于预训练的模态：except_take_put 最佳 head-only 为 0.604，take_put 最佳为 0.666。",
        "EMG 预训练有中等收益，但绝对性能较低且平均 ECE 约 0.388，过度自信明显。",
        "传统 RGB 预训练尚未稳定超过 scratch；head-only 明显退化，RGB 必须更新 backbone。",
        f"Round-2 动作语义保持增强中，{round2_best.source_key}、backbone LR={round2_best.backbone_lr:g} 达到 {round2_best.max_balanced:.3f}，相对该轮 scratch 提升 {(round2_best.max_balanced-round2_scratch.mean_balanced)*100:.1f} pp。",
        f"双相机预训练当前整体失败：最佳预训练跨相机平均 {dualcam_best.mean_balanced:.3f}，低于 scratch 的 {dualcam_scratch.mean_balanced:.3f}。",
        "首要可信度限制是仅有 seed=1、checkpoint 选择不一致，以及 RGB scratch 跨实验族波动。",
    ]
    for text in bullets: add_bullet(doc,text,bullet_id)

    doc.add_heading("1. 数据覆盖与分析规则",level=1)
    add_table(doc,["项目","数量/规则"],[
        ["CL args.json",f"原始 {inv['cl_args_files_raw']}；去重后 {inv['cl_configs_deduplicated']}；重复 canonical 配置 {inv['duplicate_cl_configs']}"],
        ["FT summary.json",str(inv['ft_training_runs'])],["测试汇总",f"{inv['test_metric_rows']} 条；best_val {inv['best_val_test_rows']}，last {inv['last_test_rows']}"],
        ["逐样本回溯",f"{inv['per_sample_available']} 个运行/相机视图可用"],
        ["分析运行环境",summary.get("analysis_runtime",{}).get("python_executable","unknown")],
        ["主视图","优先 best_val，无对应测试时使用 last；另保留严格 last 视图"],
        ["基线","优先同实验族 scratch；跨实验族比较单独标注 checkpoint 与 scratch 波动"],
    ],[1800,7560])

    doc.add_heading("2. 实验设计总览",level=1)
    add_table(doc,["实验族","研究问题","主要变量","范围"],FAMILY_DESCRIPTIONS,[1800,2100,3660,1800])

    doc.add_heading("2.1 Prototype loss 与 Relation loss 专项分析",level=2)
    doc.add_paragraph("本节把 scratch（从头训练）与 SupCon-only 都作为基线；在同一实验族、同一模态和同一微调模式内，比较 SupCon+Proto、SupCon+Rel 与 SupCon+Proto+Rel。由于当前训练仅 seed=1，配置间差值首先是描述性证据。")
    module_rows=[]
    for mod,m,bs,bu in selected_module_rows():
        module_rows.append([mod,MODULE_NAME.get(m.loss_group,m.loss_group),f"{m.family}\n{m.finetune_mode}",short(m.source_key,30),f3(m.test_balanced_acc),
                            f"{pp(bs.delta_balanced_acc)}\n[{pp(bs.ci_low)}, {pp(bs.ci_high)}]",
                            f"{pp(bu.delta_balanced_acc)}\n[{pp(bu.ci_low)}, {pp(bu.ci_high)}]"])
    add_table(doc,["模态","最佳模块","实验簇/模式","配置","Bal. Acc","Δ Scratch\n95% CI","Δ SupCon\n95% CI"],module_rows,[600,1200,1080,2460,900,1560,1560])
    doc.add_paragraph("RGB 的 SupCon+Proto 相对匹配 SupCon 提升 9.1 pp，但相对 scratch 仅提升 2.3 pp，且样本 bootstrap 区间跨 0：Proto 主要修复了较弱 SupCon 基线，并未形成对强 scratch 的明确优势。EMG 的 SupCon+Rel 与 IMU 的 K=2048 SupCon+Proto+Rel 相对两种基线均为正，区间不跨 0；但这些区间只反映固定模型的测试样本不确定性，不包括训练 seed 变异。")
    for idx,modality in enumerate(("rgb","emg","imu"),start=1):
        add_picture(doc,f"module_method_performance_{modality}.png",f"图 2.{idx}. {modality.upper()} 的损失模块比较热力图。逐行比较同一实验簇；列为 Scratch、SupCon、SupCon+Proto、SupCon+Rel、SupCon+Proto+Rel，颜色越深表示 balanced accuracy 越高；空白表示未运行。",6.15)

    doc.add_heading("3. except_take_put 主任务",level=1)
    rows=[]
    for m,s,f,h in except_rows:
        rows.append([m,f"{f3(s.test_balanced_acc)}\n{checkpoint_tag(s)}",f"{f3(f.test_balanced_acc)}\n{short(f.source_key,28)}\n{checkpoint_tag(f)}",f"{f3(h.test_balanced_acc)}\n{short(h.source_key,28)}\n{checkpoint_tag(h)}",("预训练未稳定超过强 scratch；必须 full" if m=="RGB" else "有收益；head-only 略优" if m=="EMG" else "收益最稳定；full/head 接近")])
    add_table(doc,["模态","最佳 scratch","最佳 full","最佳 head-only","判断"],rows,[720,1200,2100,2100,3240])
    add_picture(doc,"family_best_balanced_accuracy.png","图 1. except_take_put 各实验族最佳 balanced accuracy。基础/旧实验主要为 last，新扫描主要为 best_val。",6.2)
    doc.add_heading("3.1 RGB",level=2)
    doc.add_paragraph("基础损失组合中，suploss_proto_p1 的 full 微调 balanced accuracy 为 0.607，比该实验族 scratch 高 2.3 pp；但跨实验族最佳 scratch 可达 0.638。说明 prototype 预训练存在局部收益，但尚未形成可复现的跨运行优势。head-only 最佳仅 0.475，冻结 RGB backbone 会保留不适合测试域的表示。")
    doc.add_heading("3.2 EMG 与 IMU",level=2)
    doc.add_paragraph("EMG 的最佳配置来自 relation Top-k all 的 head-only（0.369）；IMU 的最佳 head-only 来自 K_queue 扫描（0.604）。IMU 对 prototype+relation 的收益更一致；EMG 配置间差异较小，且高置信错误更多。")

    doc.add_heading("4. take_put 二分类",level=1)
    rows=[[m,f3(s.test_balanced_acc),f3(f.test_balanced_acc),pp(f.test_balanced_acc-s.test_balanced_acc),f3(h.test_balanced_acc),pp(h.test_balanced_acc-s.test_balanced_acc)] for m,s,f,h in take_rows]
    add_table(doc,["模态","scratch","最佳 full","Δ scratch","最佳 head","Δ scratch"],rows,[720,1200,1440,1440,1440,3120])
    doc.add_paragraph("RGB 仅 full 微调略高于 scratch，head-only 明显退化；IMU/EMG 的 full 与 head-only 均优于 scratch。结论与主任务一致：RGB 需要更新 backbone，一维信号更适合冻结或小幅更新。")

    doc.add_heading("5. RGB 后续实验",level=1)
    add_table(doc,["实验","最佳配置","balanced accuracy","相对 scratch","判断"],[
        ["Round-2",f"{round2_best.source_key}\nblr={round2_best.backbone_lr:g}",f3(round2_best.max_balanced),pp(round2_best.max_balanced-round2_scratch.mean_balanced),"明确有效，但学习率敏感"],
        ["Dual-camera",dualcam_best.source_key,f"平均 {f3(dualcam_best.mean_balanced)}",pp(dualcam_best.mean_balanced-dualcam_scratch.mean_balanced),"所有预训练均低于 scratch"],
    ],[1200,3000,1560,1440,2160])
    add_picture(doc,"round2_and_dualcam.png","图 2. Round-2 与双相机 RGB 后续实验。双相机柱图为两个测试相机的均值。",6.25)

    doc.add_heading("6. 参数敏感性与 checkpoint",level=1)
    doc.add_paragraph("K_queue 和 Top-k 的最优点具有明显模态依赖，未出现跨 RGB/EMG/IMU 的统一最优值。参数扫描带来的平均变化通常小于模态差异，因此应优先解决表示与数据域问题。")
    add_picture(doc,"hyperparameter_sensitivity.png","图 3. 队列容量与关系损失 Top-k 的敏感性。曲线为相关配置和微调模式的平均值。",6.25)
    doc.add_paragraph("best_val 相对 last 的平均变化多数在 ±1 pp 内，但单个运行可超过 4 pp；这足以改变接近配置的排序。")
    add_picture(doc,"checkpoint_selection_effect.png","图 4. 各实验族 best_val 相对 last 的 balanced accuracy 平均变化。",6.05)

    doc.add_heading("7. 训练稳定性与校准",level=1)
    doc.add_paragraph("except_take_put 的训练-验证 accuracy gap 中位数：EMG full 约 0.238、IMU full 约 0.144、RGB full 约 0.112。IMU 与 RGB 的 head-only 通常更早达到最佳验证 epoch；EMG full/head 均在约 45 epoch 达峰。")
    cal_rows=[[r.modality.upper(),int(r.n_models),f3(r.mean_ece),f3(r.mean_confidence),f3(r.error_rate),int(r.high_conf_errors)] for _,r in cal.iterrows()]
    add_table(doc,["模态","视图数","平均 ECE","平均置信度","平均错误率","高置信错误*"],cal_rows,[840,1080,1320,1440,1440,3240])
    p=doc.add_paragraph(); set_run_font(p.add_run("注："),bold=True,color=GOLD); set_run_font(p.add_run("高置信错误为模型×样本累计次数。EMG 最严重，建议优先尝试 temperature scaling，并检查信号归一化。"),size=10)
    add_picture(doc,"selected_model_reliability.png","图 4.1. 所选 Scratch、SupCon 与最佳模块的 10-bin 可靠性图。横轴为区间平均置信度，纵轴为实际正确率；对角线表示理想校准，曲线位于其下方表示过度自信。ECE 是各区间绝对偏差按样本占比加权之和。",6.2)

    doc.add_heading("8. 类别级问题",level=1)
    add_picture(doc,"class_recall_delta_vs_scratch.png","图 5. except_take_put 最佳预训练模型相对 scratch 的逐类别 recall 变化。",6.25)
    top_conf=confusions.head(10)
    add_table(doc,["任务","模态","真实类","预测类","累计错误*"],[[r.dataset_scope,r.modality.upper(),r.true_class,r.predicted_class,int(r.error_count)] for _,r in top_conf.iterrows()],[1800,840,1680,1680,3360])
    doc.add_paragraph("RGB 预训练对 press 有明显帮助，但对 measure、label、insert 等类别可能退化；EMG 的 wrap→pull_out、insert→pull_out 需要优先回溯。累计错误受模型数量影响，用于定位问题，不用于跨模态直接比较。")

    doc.add_heading("8.1 最佳模块的逐类别 recall 变化",level=2)
    doc.add_paragraph("每个模态选择一个最佳自研模块，并与匹配的 Scratch 和 SupCon 基线比较。图的上半部分是绝对 recall，下半部分是模块减基线的变化；误差线来自逐类别配对样本 bootstrap。虚线 ±5 pp 是用于筛选实际影响的描述性阈值，不是显著性阈值。")
    for idx,modality in enumerate(("rgb","emg","imu"),start=1):
        add_picture(doc,f"selected_module_per_class_{modality}.png",f"图 5.{idx}. {modality.upper()} 最佳自研模块与 Scratch、SupCon 的逐类别 recall。下半图中正值表示模块改善该类，负值表示产生副作用；区间只反映当前固定权重的测试样本不确定性。",6.2)

    doc.add_heading("8.2 各方法对类别的作用与副作用",level=2)
    doc.add_paragraph("在基础损失实验簇内，分别为 SupCon、SupCon+Proto、SupCon+Rel、SupCon+Proto+Rel 选 balanced accuracy 最佳配置。相对 SupCon，RGB full 的 Proto 对 9 类有益、3 类有害；EMG head-only 的 Rel 对 7 类有益、1 类有害；IMU head-only 的 Proto+Rel 对 5 类有益、4 类有害。完整幅度和 source_key 已写入工作簿。")
    for idx,modality in enumerate(("rgb","emg","imu"),start=1):
        add_picture(doc,f"class_effects_{modality}.png",f"图 5.{idx+3}. {modality.upper()} 方法×类别作用矩阵。行是损失方法，列是类别，单元格为 recall 相对面板标题所示基线的变化；绿色为改善、红色为退化、近白色为小变化。",6.2)

    doc.add_heading("9. 特征分布诊断",level=1)
    doc.add_paragraph("本轮没有重新遍历全部权重，而是复用旧分析已提取的代表性 512D 特征。RGB head-only 测试 silhouette=-0.229，full=0.085；EMG 各测试表示接近或低于 0；IMU full 测试 silhouette 约 0.089。训练集与测试集分离度差距很大，域偏移仍是核心问题。")
    add_picture(doc,"feature_separation_train_test.png","图 6. 代表性配置的训练/测试特征 cosine silhouette。",6.2)
    add_picture(doc,"rgb_head_test_umap.png","图 7. RGB head-only 测试特征 UMAP：类别颜色严重混合。",5.65)
    add_picture(doc,"rgb_full_test_umap.png","图 8. RGB full 测试特征 UMAP：局部类别簇改善，但仍有重叠。",5.65)

    doc.add_heading("9.1 退化类别的定向特征空间对照",level=2)
    doc.add_paragraph("为控制计算量，本节只分析最佳模块相对 SupCon recall 确实下降的类别及其最常见混淆目标。三种模型使用同一批测试样本；特征先投影到共同 30D 子空间，再正交对齐到 SupCon，最后联合做二维 PCA。圆点表示退化类别，叉号表示其混淆目标。坐标无物理含义，应比较簇混合、紧凑度和相对邻近关系，并以原始高维指标为确认。PCA 是本轮的快速定向诊断；已有全局 UMAP 仍保留在前文。")
    diag_rows=[]
    diag = feature_diag[feature_diag.scope == "selected_harmed_and_confusion_classes"]
    for _,r in diag.iterrows():
        diag_rows.append([r.modality.upper(),r.role,int(r.n_samples),f3(r.silhouette_cosine),f3(r.knn5_purity),f3(r.within_class_cosine_distance),f3(r.nearest_other_centroid_distance),f3(r.centroid_margin)])
    add_table(doc,["模态","模型","n","Silhouette","5-NN purity","类内距离","最近类中心","中心 margin"],diag_rows,[660,900,600,1260,1200,1200,1320,1380])
    doc.add_paragraph("Silhouette 与 5-NN purity 越高通常表示局部类别分离越好；类内 cosine 距离越低表示更紧凑；最近其他类中心距离越高、中心 margin 越大表示类中心更安全。负 margin 表明平均类内扩散已超过最近类中心间距，存在明显重叠。")
    doc.add_paragraph("结果上，RGB 模块相对 SupCon 的 silhouette 略升、类内距离下降，但 5-NN purity 基本不变：整体几何略改善，却仍在 tear、insert、label 等局部边界上退化。EMG 的 silhouette 略升但类内距离增大，pull_out 更像局部边界交换。IMU 模块相对 SupCon 的 silhouette、purity 与最近类中心距离均下降，和 open recall 下降一致；因此 Proto+Rel 的总体收益并非所有类别共享。")
    for idx,modality in enumerate(("rgb","emg","imu"),start=1):
        add_picture(doc,f"harmed_class_feature_space_{modality}.png",f"图 8.{idx}. {modality.upper()} 退化类别与主要混淆目标的共享对齐 PCA。三列分别为 Scratch、SupCon 和最佳模块；比较相同颜色类别在三列中的混合与紧凑变化，不能把坐标绝对位置解释为性能。",6.2)

    doc.add_heading("10. 主要问题",level=1)
    for _,r in issues.iterrows():
        p=doc.add_paragraph(); p.paragraph_format.keep_with_next=True; set_run_font(p.add_run(f"[{r.severity}] {r.issue}"),size=11,bold=True,color=RED if r.severity=="High" else GOLD)
        p=doc.add_paragraph(); set_run_font(p.add_run("证据："),bold=True); set_run_font(p.add_run(str(r.evidence)))
        p=doc.add_paragraph(); set_run_font(p.add_run("影响："),bold=True); set_run_font(p.add_run(str(r.impact)))
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(10); set_run_font(p.add_run("建议："),bold=True,color=GREEN); set_run_font(p.add_run(str(r.recommended_action)))

    doc.add_heading("11. 推荐决策",level=1)
    recs=[
        "IMU：优先复验 K_queue/prototype+relation head-only 与基础 suploss_proto_p2 full。",
        "EMG：优先复验 relation Top-k all head-only，并加入归一化与校准消融。",
        "RGB：以 Round-2 a2_weak_shared_supcon、backbone LR=3e-4 为主线，和统一 scratch、基础 proto_p1 full 做 5-seed 对照。",
        "双相机：暂停扩大网格，先检查相机归一化、正样本定义、采样对齐以及预训练/微调域的一致性。",
        "实验治理：统一输出 best_val/last、代码 commit、manifest 哈希和确定性设置。",
    ]
    for text in recs: add_number(doc,text,number_id)

    doc.add_heading("附录 A. 严格 last 视图的最佳配置",level=1)
    rows=[]
    d=strict_last[(strict_last.dataset_scope=="except_take_put") & strict_last.finetune_mode.isin(["full","head_only","scratch"])]
    for (mod,mode),g in d.groupby(["modality","finetune_mode"]):
        r=g.sort_values("test_balanced_acc",ascending=False).iloc[0]; rows.append([mod.upper(),mode,r.family,short(r.source_key,35),f3(r.test_balanced_acc),f3(r.test_macro_f1)])
    add_table(doc,["模态","模式","实验族","配置","Bal. Acc","Macro-F1"],rows,[720,1080,1440,3240,1440,1440])

    doc.add_heading("附录 B. 指标与限制",level=1)
    for text in [
        "Balanced accuracy 是各类别 recall 的宏平均，为主排序指标。",
        "Macro-F1 是逐类别 F1 的宏平均；ECE 是 10-bin expected calibration error。",
        "Best-available 优先 best_val，无对应测试时退回 last；严格 last 视图用于同 checkpoint 比较。",
        "本报告中的提升是描述性差值，不是统计显著性结论；当前仅 seed=1。",
        "全局 UMAP/silhouette 复用旧分析的代表性特征；模块退化类别图则由本轮 9 个定向模型重新提取。两者都不代表全部 234 个 CL 配置。",
    ]: add_bullet(doc,text,bullet_id)

    doc.add_heading("附录 C. 指标计算方式、作用与限制",level=1)
    metric_rows=[[r.metric,r.calculation,r.purpose,r.interpretation_limit] for _,r in metric_defs.iterrows()]
    add_table(doc,["指标","计算方式","作用","解释限制"],metric_rows,[1500,2700,2400,2760])
    doc.add_heading("C.1 ECE 与可靠性图",level=2)
    doc.add_paragraph("将样本按最大预测概率分到 10 个置信度区间。对区间 b，acc(b) 是实际正确率，conf(b) 是平均置信度，n_b 是样本数；ECE = Σ_b (n_b/N) × |acc(b) − conf(b)|。ECE 越低说明置信度与经验正确率越一致，但它依赖分箱数和样本量，也可能掩盖类别级或局部区间问题，因此必须和可靠性曲线一起读。")
    doc.add_heading("C.2 Silhouette、邻域纯度与中心距离",level=2)
    doc.add_paragraph("Cosine silhouette 对样本 i 比较同类平均距离 a(i) 与最近其他类平均距离 b(i)：s(i)=[b(i)−a(i)]/max[a(i),b(i)]。接近 1 表示分离好，接近 0 表示边界重叠，负值表示更靠近其他类。5-NN purity 是每个样本五个最近邻中同类比例的平均；类内 cosine 距离衡量紧凑度；中心 margin=最近其他类中心距离−类内距离。所有指标都受特征归一化、距离度量、类别样本数和域偏移影响。")
    doc.add_heading("C.3 Bootstrap 区间的范围",level=2)
    doc.add_paragraph("本报告按真实类别分层，并对同一测试样本上的模型正确/错误差值做 2,000 次有放回重采样，得到 balanced accuracy 差值的 2.5% 与 97.5% 分位数。它保留了配对信息并控制类别权重，但只反映当前测试样本的抽样不确定性；由于每个配置只有 seed=1，不能覆盖重新训练带来的波动，也不是多 seed 显著性检验。")

    doc.add_heading("附录 D. 图表阅读指南",level=1)
    add_table(doc,["图表类型","视觉编码","正确阅读方式"],[[r.figure_type,r.encoding,r.how_to_read] for _,r in figure_guide.iterrows()],[1800,3780,3780])

    doc.core_properties.title="Obj2 多模态对比学习与微调实验：完整分析报告"
    doc.core_properties.subject="RGB、EMG、IMU 实验配置、性能、类别诊断与结论"
    doc.core_properties.author="Codex"
    doc.save(REPORT_DOCX)


if __name__ == "__main__":
    REPORT_MD.write_text(build_markdown(),encoding="utf-8")
    build_docx()
    print("expanded Markdown and DOCX report created")
